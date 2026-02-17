"""Document parser for extracting text from various file formats."""

import io
import logging
from typing import Optional

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

# Supported file extensions and their MIME types
SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
}


class DocumentParser:
    """Parser for extracting text from PDF, DOCX, TXT, and Markdown files."""

    @staticmethod
    def parse(file_content: bytes, filename: str) -> str:
        """Parse document and extract text content.

        Args:
            file_content: Raw file content as bytes.
            filename: Name of the file to determine format.

        Returns:
            Extracted text content.

        Raises:
            ValueError: If file format is not supported.
        """
        ext = filename.lower().split(".")[-1] if "." in filename else ""

        # Add dot if missing
        if ext and not ext.startswith("."):
            ext = f".{ext}"

        # Map extension to handler
        if ext == ".pdf":
            return DocumentParser._parse_pdf(file_content)
        elif ext in [".docx", ".doc"]:
            return DocumentParser._parse_docx(file_content)
        elif ext in [".txt", ".md", ".markdown"]:
            return DocumentParser._parse_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _parse_pdf(file_content: bytes) -> str:
        """Parse PDF file and extract text.

        Args:
            file_content: Raw PDF content.

        Returns:
            Extracted text.
        """
        try:
            with io.BytesIO(file_content) as pdf_stream:
                reader = PdfReader(pdf_stream)
                text_parts = []

                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue

                text = "\n".join(text_parts)
                logger.info(f"Extracted {len(text)} characters from PDF")
                return text

        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    @staticmethod
    def _parse_docx(file_content: bytes) -> str:
        """Parse DOCX file and extract text.

        Args:
            file_content: Raw DOCX content.

        Returns:
            Extracted text.
        """
        try:
            with io.BytesIO(file_content) as doc_stream:
                doc = Document(doc_stream)
                text_parts = []

                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text_parts.append(row_text)

                text = "\n".join(text_parts)
                logger.info(f"Extracted {len(text)} characters from DOCX")
                return text

        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def _parse_text(file_content: bytes) -> str:
        """Parse text file and extract content.

        Args:
            file_content: Raw text content.

        Returns:
            Extracted text.
        """
        # Try different encodings
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                logger.info(f"Extracted {len(text)} characters from text file (encoding: {encoding})")
                return text
            except (UnicodeDecodeError, AttributeError):
                continue

        # Fallback: decode with errors replaced
        text = file_content.decode("utf-8", errors="replace")
        logger.warning(f"Extracted text with replacement characters")
        return text

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file format is supported.

        Args:
            filename: Name of the file.

        Returns:
            True if supported, False otherwise.
        """
        ext = filename.lower().split(".")[-1] if "." in filename else ""
        if ext and not ext.startswith("."):
            ext = f".{ext}"
        return ext in SUPPORTED_EXTENSIONS
